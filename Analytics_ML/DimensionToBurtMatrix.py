# Build the Matrix of all Papers and Keys
import pandas as pd
from docx import Document
import csv
from scipy.spatial.distance import *

columns = ['D1: METH: [0|1]', 'D2: APPL [CODING]', 'D3: PROB [CODING]', 'D4: CONT [CODING]',
           'D5: AIME [CODING>EDGE|BLOC|BOTH]', 'D6: SECU [CODING>EDGE|BLOC|BOTH]',
           'D7: PRIV [CODING>EDGE|BLOC|BOTH]', 'D8: ALLO [CODING>EDGE|BLOC|BOTH]',
           'D9: METR[CODING>EDGE|BLOC|BOTH]', 'D11: TECH [CODING>EDGE|BLOC|BOTH]', 'D12: TRLE [0|1|…|9]',
           'D16: UNSD [0, 1, ...,17]', 'D17: COMM [CODING>EDGE|BLOC|BOTH]',
           'D18: EVAL [CODING]', 'D19: BLOC [CODING]', 'D20: CONS [CODING]', 'D21: PERM [0|1]',
           'D22: TYPE [0|1|2|4]', 'D23: CHAI [0/1]', 'D24: REWA [0/1/2/3]']

columns_brief = ['Methodology', 'Application', 'Problem', 'Contribution', 'AI Method', 'Security', 'Privacy',
                 'Allocation', 'Metric', 'Technology', 'TRLE', 'UNSD Code', 'Communication',
                 'Evaluation', 'Blockchain', 'Consensus', 'Permission', 'Type', 'Chain', 'Reward']


def get_feature_names(all_dimensions):
    doc = Document('Encoding.docx')
    table = doc.tables[0]
    feature_table = {}
    for d, dimension in enumerate(all_dimensions):
        features = table.cell(d + 1, 2).text.split('\n')
        feature_table[dimension] = features
    return feature_table


def handle_encoding(encoding):
    # if there is no , between ] and [ then add one
    if '][' in str(encoding):
        encoding = str(encoding).replace('][', '],[')
    if ' ' in str(encoding):
        encoding = str(encoding).replace(' ', '')
    # remove square brackets
    encoding = str(encoding).replace('[', '').replace(']', '').replace("'", "")
    # split by comma
    encoding = encoding.split(',')
    # remove spaces
    encoding = [x.strip() for x in encoding]
    # capitalize all words
    encoding = [x.upper() for x in encoding]
    return encoding


def ignore_numbers_in_encoding(encoding: list):
    encoding = [value.split('>')[0] for value in encoding]
    return encoding


def main():
    df = pd.read_excel('../ReviewedRecords.xlsx')
    for c in columns:
        df[c] = df[c].apply(lambda x: handle_encoding(x))
        df[c] = df[c].apply(lambda x: ignore_numbers_in_encoding(x))

    # explore all keywords for each dimension
    dim_num = len(columns)
    doc = Document('Encoding.docx')
    table = doc.tables[0]
    dim_table = {}
    dim_len_arr = []
    attribute_names = []
    for d in range(dim_num):
        dim_table[columns[d]] = {}
        # get all keys in each dimension
        dimension = table.cell(d + 1, 2).text.split('\n')
        # change from lower to upper, and delete the space
        dimension = [attr.upper().replace(' ', '') for attr in dimension]
        d_len = len(dimension)
        # give number of each key in the dimension for position
        values = [v for v in range(d_len)]
        d_v = dict(zip(dimension, values))
        dim_table[columns[d]].update(d_v)
        # add the number of keys in each dimension
        dim_len_arr.append(d_len)
        # write the names of attributes
        names = [columns_brief[d] + '_' + name for name in dimension]
        attribute_names = attribute_names + names
    print(dim_table)
    total_len = sum(dim_len_arr)

    # build the new matrix
    matrix = [attribute_names + ['paper_index']]
    for index, row in df.iterrows():
        # exclude the unsatisfied paper
        if not str(row["Violated [1,2,3,4]"]) == "nan":
            continue
        # generate an attribute dimension for each paper
        paper = np.zeros(shape=total_len+1)
        for i, c in enumerate(columns):
            # set the position of key in the matrix
            col_start = sum(dim_len_arr[:i])
            for item in row[c]:
                if item != "NONE" and item != "NAN":
                    if not item.isdigit():
                        # if item in dim_table:
                        if item not in dim_table[c]:
                            print(c)
                            print(index)
                            print(item)
                        else:
                            col_key_pos = col_start + dim_table[c][item]
                            paper[col_key_pos] += 1
                            paper[-1] = int(row['PaperNumber'])
                            if paper[col_key_pos] > 1:
                                paper[col_key_pos] = 1
                    else:
                        col_key_pos = col_start + int(item)
                        paper[col_key_pos] += 1
                        paper[-1] = int(row['PaperNumber'])
                        if paper[col_key_pos] > 1:
                            paper[col_key_pos] = 1

                # else:
                #     # None: add 1 at the end of each key's position
                #     none_pos = col_start + dim_len_arr[i] - 1
                #     paper[none_pos] += 1
        matrix.append(paper)
    print(matrix[-3])
    print(np.sum(matrix[-3]))
    print("The length of dimension is " + str(total_len))
    print("The number of papers is " + str(len(matrix)))

    # Store the generated matrix
    output_file = 'output/binary_matrix.csv'
    with open(output_file, 'w', newline='', encoding='utf-8-sig') as outFile:
        csv_writer = csv.writer(outFile)
        for line in matrix:
            csv_writer.writerow(line)
    outFile.close()


if __name__ == '__main__':
    main()
